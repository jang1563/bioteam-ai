"""Manuscript Studio API.

Persistent orchestration layer that aggregates linked W11, W8, W7, and
RCMXT-bearing workflow outputs into one manuscript-defense session.
"""

from __future__ import annotations

import base64
import io
import re
import tempfile
from datetime import datetime, timezone
from html import escape as _esc
from pathlib import Path
from typing import Literal
from uuid import uuid4

from app.db.database import engine as db_engine
from app.models.agent import AgentOutput
from app.models.integrity import AuditFinding, AuditRun
from app.models.manuscript import (
    IntegrityAuditReport,
    ManuscriptClaimEvidence,
    ManuscriptFallbackFlag,
    ManuscriptIntegrityFlag,
    ManuscriptOutlineSection,
    ManuscriptReportRunMetadata,
    ManuscriptReviewerRisk,
    ManuscriptSession,
    ReviewerRiskReport,
)
from app.models.story_frame import StoryFrame
from app.models.workflow import WorkflowInstance
from fastapi import APIRouter, File, HTTPException, UploadFile
from pydantic import BaseModel, Field, ValidationError
from sqlmodel import Session, select

router = APIRouter(prefix="/api/v1/manuscript", tags=["manuscript"])

_SUPPORTED_LINK_TEMPLATES = {"W11", "W8", "W7", "W1", "W6"}
_REVIEWER_UPLOAD_SUFFIXES = {".pdf", ".docx", ".doc"}
_auditor_agent = None


def set_auditor_agent(agent) -> None:
    """Wire DataIntegrityAuditorAgent during app startup."""
    global _auditor_agent
    _auditor_agent = agent


class CreateManuscriptSessionRequest(BaseModel):
    """Create a new manuscript session."""

    title: str | None = Field(default=None, max_length=200)
    query: str = Field(min_length=1, max_length=2000)
    notes: str = Field(default="", max_length=20000)
    draft_text: str = Field(default="", max_length=50000)
    target_journal: str | None = Field(default=None, max_length=200)
    key_papers: list[str] = Field(default_factory=list)


class LinkWorkflowRequest(BaseModel):
    """Link an existing workflow run into a manuscript session."""

    workflow_id: str = Field(min_length=1, max_length=100)


class SelectStoryFrameRequest(BaseModel):
    """Select one story frame inside the aggregated manuscript session."""

    frame_id: str = Field(min_length=1, max_length=100)


class ApproveStoryScopeRequest(BaseModel):
    """Approve the W11 scope checkpoint and choose a target journal tier."""

    target_tier: Literal["nature_cell", "specialty", "grant"] = "specialty"


class RunReviewerRisksRequest(BaseModel):
    """Start a W8 reviewer-risk run from a manuscript session."""

    pdf_path: str = Field(min_length=1, max_length=500)


class ManuscriptStageStatus(BaseModel):
    """Computed status for one manuscript-defense stage."""

    stage: str
    status: Literal["not_started", "running", "waiting_human", "ready", "failed", "partial"]
    detail: str
    source_workflow_id: str | None = None


class ManuscriptSessionResponse(BaseModel):
    """API response for a manuscript session."""

    id: str
    title: str
    query: str
    notes: str = ""
    draft_text: str = ""
    target_journal: str | None = None
    key_papers: list[str] = Field(default_factory=list)
    phase: str
    selected_frame_id: str | None = None
    completion_state: str
    linked_workflows: dict[str, str] = Field(default_factory=dict)
    stage_statuses: list[ManuscriptStageStatus] = Field(default_factory=list)
    fallback_flags: list[ManuscriptFallbackFlag] = Field(default_factory=list)
    frame_options: list[StoryFrame] = Field(default_factory=list)
    selected_frame: StoryFrame | None = None
    claim_map: list[ManuscriptClaimEvidence] = Field(default_factory=list)
    reviewer_risks: list[ManuscriptReviewerRisk] = Field(default_factory=list)
    reviewer_risk_report: ReviewerRiskReport | None = None
    integrity_flags: list[ManuscriptIntegrityFlag] = Field(default_factory=list)
    integrity_audit_report: IntegrityAuditReport | None = None
    outline: list[ManuscriptOutlineSection] = Field(default_factory=list)
    created_at: str
    updated_at: str


class ManuscriptDefenseBriefResponse(BaseModel):
    """Exportable manuscript defense brief."""

    session_id: str
    title: str
    filename: str
    markdown: str


class ManuscriptDefenseBriefPrintResponse(BaseModel):
    """Print-friendly manuscript defense brief."""

    session_id: str
    title: str
    filename: str
    html: str


class ManuscriptDefenseBriefDocxResponse(BaseModel):
    """Downloadable DOCX manuscript defense brief."""

    session_id: str
    title: str
    filename: str
    mime_type: str = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    content_base64: str


def _get_session_or_404(manuscript_session_id: str) -> ManuscriptSession:
    with Session(db_engine) as session:
        row = session.get(ManuscriptSession, manuscript_session_id)
        if row is None:
            raise HTTPException(status_code=404, detail=f"Manuscript session not found: {manuscript_session_id}")
        session.expunge(row)
        return row


def _save_session(row: ManuscriptSession) -> ManuscriptSession:
    row.updated_at = datetime.now(timezone.utc)
    with Session(db_engine) as session:
        merged = session.merge(row)
        session.commit()
        session.refresh(merged)
        session.expunge(merged)
        return merged


def _get_workflow_or_404(workflow_id: str) -> WorkflowInstance:
    with Session(db_engine) as session:
        row = session.get(WorkflowInstance, workflow_id)
        if row is None:
            raise HTTPException(status_code=404, detail=f"Workflow not found: {workflow_id}")
        session.expunge(row)
        return row


def _axis_summary(score: dict) -> str:
    pieces = []
    for axis in ("R", "C", "M", "X", "T"):
        value = score.get(axis)
        if value is None:
            continue
        pieces.append(f"{axis} {value:.2f}" if isinstance(value, (float, int)) else f"{axis} {value}")
    composite = score.get("composite")
    if composite is not None:
        pieces.append(f"Composite {composite:.2f}" if isinstance(composite, (float, int)) else f"Composite {composite}")
    return " · ".join(pieces)


def _claim_risk_level(composite: float | None) -> Literal["low", "medium", "high"]:
    if composite is None:
        return "medium"
    if composite < 0.45:
        return "high"
    if composite < 0.7:
        return "medium"
    return "low"


def _reviewer_risk_level(category: str) -> Literal["low", "medium", "high"]:
    if category == "major":
        return "high"
    if category == "minor":
        return "medium"
    return "low"


def _parse_story_frames(w11_workflow: WorkflowInstance | None) -> tuple[list[dict], dict | None, list[dict]]:
    if w11_workflow is None:
        return [], None, []

    manifest = dict(w11_workflow.session_manifest or {})
    frame_blob = manifest.get("frame_options") or {}
    frames = frame_blob.get("frames", []) if isinstance(frame_blob, dict) else []
    selected = manifest.get("selected_story_frame")

    fallback_flags: list[dict] = []
    generation_mode = frame_blob.get("generation_mode") if isinstance(frame_blob, dict) else None
    fallback_reason = frame_blob.get("fallback_reason") if isinstance(frame_blob, dict) else None
    if generation_mode == "synthetic_fallback":
        fallback_flags.append(
            ManuscriptFallbackFlag(
                stage="story_frames",
                detail=fallback_reason or "Story frames were generated via synthetic fallback.",
                source_workflow_id=w11_workflow.id,
                provenance="synthetic_fallback",
            ).model_dump(mode="json")
        )

    if isinstance(selected, dict) and selected.get("provenance") == "synthetic_fallback":
        fallback_flags.append(
            ManuscriptFallbackFlag(
                stage="story_frames",
                detail="Selected frame was produced by the synthetic fallback path.",
                source_workflow_id=w11_workflow.id,
                provenance="synthetic_fallback",
            ).model_dump(mode="json")
        )

    return frames, selected if isinstance(selected, dict) else None, fallback_flags


def _load_w8_report(w8_workflow: WorkflowInstance | None) -> tuple[dict | None, list[dict]]:
    if w8_workflow is None:
        return None, []

    manifest = dict(w8_workflow.session_manifest or {})
    report = manifest.get("peer_review_report")
    fallback_flags: list[dict] = []
    if not isinstance(report, dict):
        fallback_flags.append(
            ManuscriptFallbackFlag(
                stage="reviewer_risks",
                detail="This W8 workflow does not have a persisted peer review report yet. Link a newer run to surface concern-level reviewer risks.",
                source_workflow_id=w8_workflow.id,
                provenance="partial_persistence",
            ).model_dump(mode="json")
        )
        return None, fallback_flags
    return report, fallback_flags


def _build_claim_map(workflows: dict[str, WorkflowInstance], w8_report: dict | None) -> list[dict]:
    claims_lookup: dict[str, dict] = {}
    if isinstance(w8_report, dict):
        for item in w8_report.get("claims_extracted", []) or []:
            if isinstance(item, dict) and item.get("claim_text"):
                claims_lookup[item["claim_text"]] = item

    items: list[dict] = []
    seen: set[str] = set()
    for template in ("W8", "W1", "W6"):
        workflow = workflows.get(template)
        if workflow is None:
            continue
        for raw_score in workflow.rcmxt_scores or []:
            claim_text = raw_score.get("claim", "").strip()
            if not claim_text or claim_text in seen:
                continue
            seen.add(claim_text)
            extracted = claims_lookup.get(claim_text, {})
            supporting_sources = list(raw_score.get("sources", []) or [])
            supporting_sources.extend(extracted.get("supporting_refs", []) or [])
            deduped_sources = list(dict.fromkeys(supporting_sources))
            item = ManuscriptClaimEvidence(
                claim_text=claim_text,
                composite_score=raw_score.get("composite"),
                axis_scores={axis: raw_score.get(axis) for axis in ("R", "C", "M", "X", "T")},
                supporting_sources=deduped_sources,
                rcmxt_summary=_axis_summary(raw_score),
                risk_level=_claim_risk_level(raw_score.get("composite")),
                generated_by=f"{template}:{workflow.id}",
            )
            items.append(item.model_dump(mode="json"))

    items.sort(
        key=lambda item: (
            {"high": 0, "medium": 1, "low": 2}.get(item.get("risk_level", "medium"), 1),
            -(item.get("composite_score") or 0),
        )
    )
    return items


def _build_reviewer_risks(w8_workflow: WorkflowInstance | None, w8_report: dict | None) -> list[dict]:
    if w8_workflow is None or not isinstance(w8_report, dict):
        return []

    generated_by = f"W8:{w8_workflow.id}"
    risks: list[dict] = []

    synthesis = w8_report.get("synthesis") or {}
    for comment in synthesis.get("comments", []) or []:
        if not isinstance(comment, dict):
            continue
        text = str(comment.get("comment", "")).strip()
        if not text:
            continue
        risks.append(
            ManuscriptReviewerRisk(
                title=text[:160],
                severity=_reviewer_risk_level(str(comment.get("category", "minor"))),
                section=str(comment.get("section", "General")),
                detail=text,
                evidence_basis=str(comment.get("evidence_basis", "")),
                generated_by=generated_by,
            ).model_dump(mode="json")
        )

    novelty = w8_report.get("novelty_assessment") or {}
    for missing in novelty.get("landmark_papers_missing", []) or []:
        risks.append(
            ManuscriptReviewerRisk(
                title="Missing landmark comparison",
                severity="medium",
                section="Novelty",
                detail=str(missing),
                generated_by=generated_by,
            ).model_dump(mode="json")
        )

    methodology = w8_report.get("methodology_assessment") or {}
    for issue in methodology.get("potential_biases", []) or []:
        risks.append(
            ManuscriptReviewerRisk(
                title="Potential methodological bias",
                severity="high",
                section="Methods",
                detail=str(issue),
                generated_by=generated_by,
            ).model_dump(mode="json")
        )
    for issue in methodology.get("reproducibility_concerns", []) or []:
        risks.append(
            ManuscriptReviewerRisk(
                title="Reproducibility concern",
                severity="high",
                section="Methods",
                detail=str(issue),
                generated_by=generated_by,
            ).model_dump(mode="json")
        )

    return risks[:12]


def _build_integrity_flags(w7_workflow: WorkflowInstance | None) -> list[dict]:
    if w7_workflow is None:
        return []

    generated_by = f"W7:{w7_workflow.id}"
    with Session(db_engine) as session:
        finding_rows = session.exec(
            select(AuditFinding)
            .where(AuditFinding.workflow_id == w7_workflow.id)
            .order_by(AuditFinding.created_at.desc())
        ).all()
        for row in finding_rows:
            session.expunge(row)

    if finding_rows:
        return [
            ManuscriptIntegrityFlag(
                title=row.title or row.category.replace("_", " ").title(),
                severity=row.severity,
                category=row.category,
                detail=row.description or row.source_text,
                suggestion=row.suggestion,
                status=row.status,
                generated_by=generated_by,
            ).model_dump(mode="json")
            for row in finding_rows[:12]
        ]

    report = (w7_workflow.session_manifest or {}).get("integrity_report", {})
    findings = report.get("findings", []) if isinstance(report, dict) else []
    return [
        ManuscriptIntegrityFlag(
            title=str(item.get("title") or item.get("category", "Integrity finding")).strip(),
            severity=str(item.get("severity", "warning")),
            category=str(item.get("category", "unknown")),
            detail=str(item.get("description") or item.get("source_text", "")).strip(),
            suggestion=str(item.get("suggestion", "")).strip(),
            status="open",
            generated_by=generated_by,
        ).model_dump(mode="json")
        for item in findings[:12]
        if isinstance(item, dict)
    ]


def _workflow_report_metadata(
    workflow: WorkflowInstance | None,
    *,
    template: str,
    source_count: int,
) -> ManuscriptReportRunMetadata:
    return ManuscriptReportRunMetadata(
        workflow_id=workflow.id if workflow else None,
        workflow_template=template,
        workflow_state=workflow.state if workflow else "",
        generated_at=workflow.updated_at if workflow else None,
        report_version="v1",
        source_count=source_count,
    )


def _build_reviewer_risk_report(
    w8_workflow: WorkflowInstance | None,
    w8_report: dict | None,
    reviewer_risks: list[dict],
) -> ReviewerRiskReport | None:
    if w8_workflow is None or not isinstance(w8_report, dict):
        return None

    synthesis = w8_report.get("synthesis") or {}
    claims = w8_report.get("claims_extracted") or []
    novelty = w8_report.get("novelty_assessment") or {}
    methodology = w8_report.get("methodology_assessment") or {}

    evidence_provenance: list[str] = []
    if claims:
        evidence_provenance.append(f"Claim extraction grounded {len(claims)} manuscript claims for review synthesis.")
    missing_landmarks = novelty.get("landmark_papers_missing") or []
    if missing_landmarks:
        evidence_provenance.append(
            f"Novelty assessment surfaced {len(missing_landmarks)} missing landmark comparisons."
        )
    if methodology:
        evidence_provenance.append(
            "Methodology assessment covered controls, sample size, reproducibility, and bias checks."
        )
    if not evidence_provenance:
        evidence_provenance.append("Reviewer risks were derived from the linked W8 peer-review report.")

    confidence = synthesis.get("confidence_in_conclusions")
    confidence_or_coverage = f"{len(reviewer_risks)} surfaced reviewer risks from the linked W8 run."
    if isinstance(confidence, (float, int)):
        confidence_or_coverage += f" Confidence in conclusions: {confidence:.2f}."

    summary = str(synthesis.get("summary_assessment", "")).strip()
    if not summary:
        summary = f"{len(reviewer_risks)} reviewer risks were surfaced from the linked W8 analysis."

    return ReviewerRiskReport(
        summary=summary,
        findings=[ManuscriptReviewerRisk(**item) for item in reviewer_risks if isinstance(item, dict)],
        evidence_provenance=evidence_provenance,
        confidence_or_coverage=confidence_or_coverage,
        run_metadata=_workflow_report_metadata(
            w8_workflow,
            template="W8",
            source_count=len(reviewer_risks),
        ),
    )


def _build_integrity_audit_report(
    w7_workflow: WorkflowInstance | None,
    integrity_flags: list[dict],
) -> IntegrityAuditReport | None:
    if not integrity_flags and w7_workflow is None:
        return None

    quick_scan_count = sum(
        1
        for item in integrity_flags
        if isinstance(item, dict) and str(item.get("generated_by", "")).startswith("manuscript_session:quick_check")
    )
    full_audit_count = sum(
        1 for item in integrity_flags if isinstance(item, dict) and str(item.get("generated_by", "")).startswith("W7:")
    )

    report_manifest = dict((w7_workflow.session_manifest or {}).get("integrity_report", {})) if w7_workflow else {}
    audit_run: AuditRun | None = None
    if w7_workflow is not None:
        with Session(db_engine) as session:
            audit_run = session.exec(
                select(AuditRun)
                .where(AuditRun.workflow_id == w7_workflow.id)
                .order_by(AuditRun.created_at.desc())
            ).first()
            if audit_run is not None:
                session.expunge(audit_run)

    evidence_provenance: list[str] = []
    if full_audit_count:
        evidence_provenance.append(f"Full W7 audit contributed {full_audit_count} persisted findings.")
    if quick_scan_count:
        evidence_provenance.append(f"Quick manuscript draft scan contributed {quick_scan_count} findings.")
    findings_by_category = report_manifest.get("findings_by_category")
    if isinstance(findings_by_category, dict) and findings_by_category:
        category_summary = ", ".join(
            f"{str(category).replace('_', ' ')} ({count})"
            for category, count in list(findings_by_category.items())[:4]
        )
        evidence_provenance.append(f"Category coverage: {category_summary}.")
    if not evidence_provenance:
        evidence_provenance.append("Submission checks were derived from the linked manuscript-session findings.")

    summary = ""
    if audit_run is not None and audit_run.summary:
        summary = audit_run.summary.strip()
    if not summary:
        summary = str(report_manifest.get("summary", "")).strip()
    if not summary:
        overall_level = str(report_manifest.get("overall_level", "")).strip().replace("_", " ")
        summary = f"{len(integrity_flags)} submission-check findings are currently attached to this manuscript session."
        if overall_level:
            summary += f" Overall level: {overall_level}."

    coverage_parts: list[str] = []
    if audit_run is not None and audit_run.trigger:
        coverage_parts.append(f"Trigger: {audit_run.trigger}")
    if full_audit_count:
        coverage_parts.append(f"full audit findings {full_audit_count}")
    if quick_scan_count:
        coverage_parts.append(f"quick scan findings {quick_scan_count}")
    confidence_or_coverage = "; ".join(coverage_parts) if coverage_parts else "No audit-coverage metadata available."

    return IntegrityAuditReport(
        maturity="validated_core" if w7_workflow else "guided_support",
        summary=summary,
        findings=[ManuscriptIntegrityFlag(**item) for item in integrity_flags if isinstance(item, dict)],
        evidence_provenance=evidence_provenance,
        confidence_or_coverage=confidence_or_coverage,
        run_metadata=_workflow_report_metadata(
            w7_workflow,
            template="W7",
            source_count=len(integrity_flags),
        ),
    )


def _build_outline(selected_frame: dict | None, claim_map: list[dict], reviewer_risks: list[dict], integrity_flags: list[dict]) -> list[dict]:
    if not isinstance(selected_frame, dict):
        return []

    sections: list[dict] = []
    sections.append(
        ManuscriptOutlineSection(
            title="Narrative Anchor",
            bullets=[
                str(selected_frame.get("hook", "")).strip(),
                f"Core claim: {selected_frame.get('core_claim', '')}".strip(),
                f"Central tension: {selected_frame.get('central_tension', '')}".strip(),
            ],
            generated_by="manuscript_session:derived",
        ).model_dump(mode="json")
    )

    supporting = [str(item).strip() for item in selected_frame.get("supporting_findings", []) if str(item).strip()]
    if not supporting:
        supporting = [
            f"{item.get('claim_text', '')} ({item.get('rcmxt_summary', '')})"
            for item in claim_map[:3]
            if item.get("claim_text")
        ]
    if supporting:
        sections.append(
            ManuscriptOutlineSection(
                title="Evidence to Emphasize",
                bullets=supporting[:5],
                generated_by="manuscript_session:derived",
            ).model_dump(mode="json")
        )

    figure_sequence = [str(item).strip() for item in selected_frame.get("figure_sequence", []) if str(item).strip()]
    if figure_sequence:
        sections.append(
            ManuscriptOutlineSection(
                title="Figure and Flow",
                bullets=figure_sequence[:6],
                generated_by="manuscript_session:derived",
            ).model_dump(mode="json")
        )

    defense_bullets: list[str] = []
    for risk in reviewer_risks[:2]:
        defense_bullets.append(f"Reviewer risk: {risk.get('detail') or risk.get('title')}")
    for flag in integrity_flags[:2]:
        defense_bullets.append(f"Submission check: {flag.get('title')} ({flag.get('severity')})")
    if defense_bullets:
        sections.append(
            ManuscriptOutlineSection(
                title="Defense Notes",
                bullets=defense_bullets,
                generated_by="manuscript_session:derived",
            ).model_dump(mode="json")
        )

    return sections


def _merge_unique_items(existing: list[dict], incoming: list[dict], keys: tuple[str, ...]) -> list[dict]:
    merged: list[dict] = []
    seen: set[tuple[str, ...]] = set()
    for item in [*existing, *incoming]:
        if not isinstance(item, dict):
            continue
        marker = tuple(str(item.get(key, "")) for key in keys)
        if marker in seen:
            continue
        seen.add(marker)
        merged.append(item)
    return merged


def _build_stage_statuses(
    row: ManuscriptSession,
    workflows: dict[str, WorkflowInstance],
    claim_map: list[dict],
    reviewer_risks: list[dict],
    integrity_flags: list[dict],
) -> list[dict]:
    statuses: list[dict] = []

    w11 = workflows.get("W11")
    if w11 is None:
        statuses.append(
            ManuscriptStageStatus(
                stage="story_frames",
                status="not_started",
                detail="Start or link a W11 workflow to generate narrative frames.",
            ).model_dump(mode="json")
        )
    else:
        manifest = dict(w11.session_manifest or {})
        phase = str(manifest.get("w11_phase", "")).strip().lower()
        selection_error = str(manifest.get("selection_error", "")).strip()
        scope_summary = str(manifest.get("scope_summary", "")).strip()

        if w11.state == "WAITING_HUMAN":
            if phase == "awaiting_scope":
                detail = scope_summary or "W11 is waiting for scope approval and target tier selection."
            elif phase == "awaiting_selection":
                detail = selection_error or "W11 is waiting for you to choose a story frame."
            else:
                detail = "W11 is waiting for scope approval or frame selection."
            statuses.append(
                ManuscriptStageStatus(
                    stage="story_frames",
                    status="waiting_human",
                    detail=detail,
                    source_workflow_id=w11.id,
                ).model_dump(mode="json")
            )
        elif w11.state in {"RUNNING", "PENDING"}:
            if phase == "awaiting_selection" and row.selected_frame_id:
                detail = "W11 is finalizing the selected story frame."
            elif phase == "awaiting_selection":
                detail = "W11 is preparing story frames for your review."
            else:
                detail = "W11 is generating or preparing story frames."
            statuses.append(
                ManuscriptStageStatus(
                    stage="story_frames",
                    status="running",
                    detail=detail,
                    source_workflow_id=w11.id,
                ).model_dump(mode="json")
            )
        elif w11.state in {"FAILED", "CANCELLED", "OVER_BUDGET"}:
            statuses.append(
                ManuscriptStageStatus(
                    stage="story_frames",
                    status="failed",
                    detail=f"W11 ended in state {w11.state}.",
                    source_workflow_id=w11.id,
                ).model_dump(mode="json")
            )
        elif row.selected_frame_id or row.selected_frame:
            statuses.append(
                ManuscriptStageStatus(
                    stage="story_frames",
                    status="ready",
                    detail="A story frame has been selected for this manuscript session.",
                    source_workflow_id=w11.id,
                ).model_dump(mode="json")
            )
        else:
            statuses.append(
                ManuscriptStageStatus(
                    stage="story_frames",
                    status="partial",
                    detail="W11 completed, but the session does not have a selected story frame yet.",
                    source_workflow_id=w11.id,
                ).model_dump(mode="json")
            )

    if claim_map:
        statuses.append(
            ManuscriptStageStatus(
                stage="claim_map",
                status="ready",
                detail=f"{len(claim_map)} claims have evidence calibration attached.",
                source_workflow_id=(workflows.get("W8") or workflows.get("W1") or workflows.get("W6")).id if (workflows.get("W8") or workflows.get("W1") or workflows.get("W6")) else None,
            ).model_dump(mode="json")
        )
    else:
        statuses.append(
            ManuscriptStageStatus(
                stage="claim_map",
                status="not_started",
                detail="Link a workflow with RCMXT scores to build the claim map.",
            ).model_dump(mode="json")
        )

    w8 = workflows.get("W8")
    if reviewer_risks:
        statuses.append(
            ManuscriptStageStatus(
                stage="reviewer_risks",
                status="ready",
                detail=f"{len(reviewer_risks)} reviewer risks are available.",
                source_workflow_id=w8.id if w8 else None,
            ).model_dump(mode="json")
        )
    elif w8 is None:
        statuses.append(
            ManuscriptStageStatus(
                stage="reviewer_risks",
                status="not_started",
                detail="Link a W8 workflow with a reviewed paper to surface reviewer risks.",
            ).model_dump(mode="json")
        )
    elif w8.state == "WAITING_HUMAN":
        statuses.append(
            ManuscriptStageStatus(
                stage="reviewer_risks",
                status="waiting_human",
                detail="W8 is paused at a reviewer checkpoint.",
                source_workflow_id=w8.id,
            ).model_dump(mode="json")
        )
    elif w8.state in {"RUNNING", "PENDING"}:
        statuses.append(
            ManuscriptStageStatus(
                stage="reviewer_risks",
                status="running",
                detail="W8 is analyzing the paper and assembling review risks.",
                source_workflow_id=w8.id,
            ).model_dump(mode="json")
        )
    else:
        statuses.append(
            ManuscriptStageStatus(
                stage="reviewer_risks",
                status="partial",
                detail="A W8 workflow is linked, but concern-level risks are not yet available.",
                source_workflow_id=w8.id,
            ).model_dump(mode="json")
        )

    w7 = workflows.get("W7")
    manual_check = any(
        isinstance(item, dict) and str(item.get("generated_by", "")).startswith("manuscript_session:quick_check")
        for item in integrity_flags
    )
    if w7 and w7.state in {"RUNNING", "PENDING"}:
        statuses.append(
            ManuscriptStageStatus(
                stage="submission_checks",
                status="running",
                detail=(
                    "W7 is running a full submission audit. Quick draft scan findings remain visible while the full audit is in progress."
                    if manual_check
                    else "W7 is currently evaluating submission issues."
                ),
                source_workflow_id=w7.id,
            ).model_dump(mode="json")
        )
    elif w7 and w7.state in {"FAILED", "CANCELLED", "OVER_BUDGET"}:
        statuses.append(
            ManuscriptStageStatus(
                stage="submission_checks",
                status="failed",
                detail=(
                    f"W7 ended in state {w7.state}. Quick draft scan findings are still available."
                    if manual_check
                    else f"W7 ended in state {w7.state}."
                ),
                source_workflow_id=w7.id,
            ).model_dump(mode="json")
        )
    elif integrity_flags:
        statuses.append(
            ManuscriptStageStatus(
                stage="submission_checks",
                status="ready",
                detail=(
                    "Submission checks include full W7 audit findings plus the quick draft scan."
                    if w7 and manual_check
                    else "Submission checks are populated from linked W7 results."
                    if w7
                    else "Submission checks were generated from the quick draft scan on session text."
                ),
                source_workflow_id=w7.id if w7 else None,
            ).model_dump(mode="json")
        )
    elif w7 is None and not manual_check:
        statuses.append(
            ManuscriptStageStatus(
                stage="submission_checks",
                status="not_started",
                detail="Run a quick draft scan from the session text or start a full W7 submission audit.",
            ).model_dump(mode="json")
        )

    return statuses


def _derive_phase(frame_options: list[dict], selected_frame: dict | None, claim_map: list[dict], reviewer_risks: list[dict], integrity_flags: list[dict]) -> tuple[str, str]:
    if not frame_options:
        return "collect_inputs", "empty"
    if not selected_frame:
        return "select_frame", "partial"
    if not claim_map:
        return "review_claims", "partial"
    if not reviewer_risks and not integrity_flags:
        return "defense_checks", "partial"
    return "outline_ready", "ready"


def _sync_session_outputs(row: ManuscriptSession) -> ManuscriptSession:
    linked_workflows = dict(row.linked_workflows or {})

    workflows: dict[str, WorkflowInstance] = {}
    missing_workflow_flags: list[dict] = []
    with Session(db_engine) as session:
        for template, workflow_id in linked_workflows.items():
            workflow = session.get(WorkflowInstance, workflow_id)
            if workflow is None:
                missing_workflow_flags.append(
                    ManuscriptFallbackFlag(
                        stage="linked_workflow",
                        detail=f"Linked workflow {workflow_id} for {template} is no longer available.",
                        source_workflow_id=workflow_id,
                        provenance="missing_workflow",
                    ).model_dump(mode="json")
                )
                continue
            session.expunge(workflow)
            workflows[template] = workflow

    frame_options, workflow_selected_frame, frame_fallback_flags = _parse_story_frames(workflows.get("W11"))
    w8_report, w8_fallback_flags = _load_w8_report(workflows.get("W8"))

    existing_integrity = [
        item
        for item in (row.integrity_flags or [])
        if isinstance(item, dict) and str(item.get("generated_by", "")).startswith("manuscript_session:quick_check")
    ]
    existing_fallbacks = [
        item
        for item in (row.fallback_flags or [])
        if isinstance(item, dict) and item.get("stage") == "submission_checks"
    ]

    selected_frame = workflow_selected_frame
    if row.selected_frame_id and frame_options:
        selected_frame = next(
            (frame for frame in frame_options if str(frame.get("frame_id", "")).upper() == row.selected_frame_id.upper()),
            selected_frame,
        )
    elif isinstance(workflow_selected_frame, dict):
        row.selected_frame_id = str(workflow_selected_frame.get("frame_id", "")) or None

    claim_map = _build_claim_map(workflows, w8_report)
    reviewer_risks = _build_reviewer_risks(workflows.get("W8"), w8_report)
    integrity_flags = _build_integrity_flags(workflows.get("W7"))
    integrity_flags = _merge_unique_items(existing_integrity, integrity_flags, ("title", "detail", "generated_by"))
    outline = _build_outline(selected_frame, claim_map, reviewer_risks, integrity_flags)
    phase, completion_state = _derive_phase(frame_options, selected_frame, claim_map, reviewer_risks, integrity_flags)
    row.frame_options = frame_options
    row.selected_frame = selected_frame
    row.claim_map = claim_map
    row.reviewer_risks = reviewer_risks
    row.integrity_flags = integrity_flags
    row.outline = outline
    row.phase = phase
    row.completion_state = completion_state
    row.fallback_flags = _merge_unique_items(
        existing_fallbacks,
        frame_fallback_flags + w8_fallback_flags + missing_workflow_flags,
        ("stage", "detail", "source_workflow_id"),
    )
    return row


def _load_linked_workflows(row: ManuscriptSession) -> dict[str, WorkflowInstance]:
    linked_workflows = dict(row.linked_workflows or {})
    workflows: dict[str, WorkflowInstance] = {}
    with Session(db_engine) as session:
        for template, workflow_id in linked_workflows.items():
            workflow = session.get(WorkflowInstance, workflow_id)
            if workflow is None:
                continue
            session.expunge(workflow)
            workflows[template] = workflow
    return workflows


def _to_response(row: ManuscriptSession) -> ManuscriptSessionResponse:
    workflows = _load_linked_workflows(row)
    w8_workflow = workflows.get("W8")
    w8_report = _load_w8_report(w8_workflow)[0]
    stage_statuses = _build_stage_statuses(
        row,
        workflows,
        list(row.claim_map or []),
        list(row.reviewer_risks or []),
        list(row.integrity_flags or []),
    )
    return ManuscriptSessionResponse(
        id=row.id,
        title=row.title,
        query=row.query,
        notes=row.notes,
        draft_text=row.draft_text,
        target_journal=row.target_journal,
        key_papers=list(row.key_papers or []),
        phase=row.phase,
        selected_frame_id=row.selected_frame_id,
        completion_state=row.completion_state,
        linked_workflows=dict(row.linked_workflows or {}),
        stage_statuses=[ManuscriptStageStatus(**item) for item in stage_statuses if isinstance(item, dict)],
        fallback_flags=[ManuscriptFallbackFlag(**item) for item in row.fallback_flags or [] if isinstance(item, dict)],
        frame_options=[StoryFrame(**item) for item in row.frame_options or [] if isinstance(item, dict)],
        selected_frame=StoryFrame(**row.selected_frame) if isinstance(row.selected_frame, dict) else None,
        claim_map=[ManuscriptClaimEvidence(**item) for item in row.claim_map or [] if isinstance(item, dict)],
        reviewer_risks=[ManuscriptReviewerRisk(**item) for item in row.reviewer_risks or [] if isinstance(item, dict)],
        reviewer_risk_report=_build_reviewer_risk_report(w8_workflow, w8_report, list(row.reviewer_risks or [])),
        integrity_flags=[ManuscriptIntegrityFlag(**item) for item in row.integrity_flags or [] if isinstance(item, dict)],
        integrity_audit_report=_build_integrity_audit_report(workflows.get("W7"), list(row.integrity_flags or [])),
        outline=[ManuscriptOutlineSection(**item) for item in row.outline or [] if isinstance(item, dict)],
        created_at=row.created_at.isoformat(),
        updated_at=row.updated_at.isoformat(),
    )


def _markdown_list(items: list[str]) -> list[str]:
    return [f"- {item}" for item in items if item]


def _export_filename(row: ManuscriptSession, extension: str = "md") -> str:
    slug = re.sub(r"[^A-Za-z0-9._-]+", "_", row.title or "manuscript_defense_brief").strip("._")
    if not slug:
        slug = "manuscript_defense_brief"
    return f"{slug.lower()}_defense_brief.{extension}"


def _build_defense_brief_markdown(row: ManuscriptSession) -> str:
    lines: list[str] = [
        f"# {row.title}",
        "",
        "## Manuscript Objective",
        row.query.strip() or "No manuscript objective provided.",
        "",
        "## Session Status",
        f"- Phase: {row.phase}",
        f"- Completion: {row.completion_state}",
    ]
    if row.target_journal:
        lines.append(f"- Target journal: {row.target_journal}")
    if row.linked_workflows:
        lines.extend(
            _markdown_list(
                [f"Linked {template}: {workflow_id}" for template, workflow_id in sorted(row.linked_workflows.items())]
            )
        )
    if row.notes.strip():
        lines.extend(["", "## Notes", row.notes.strip()])

    lines.extend(["", "## Story Frame"])
    if isinstance(row.selected_frame, dict):
        selected_frame = row.selected_frame
        lines.extend(
            _markdown_list(
                [
                    f"Selected frame: {selected_frame.get('frame_id', '')}",
                    f"Narrative type: {str(selected_frame.get('narrative_type', '')).replace('_', ' ')}",
                    f"Target tier: {selected_frame.get('target_tier', '')}",
                    f"Hook: {selected_frame.get('hook', '')}",
                    f"Core claim: {selected_frame.get('core_claim', '')}",
                    f"Central tension: {selected_frame.get('central_tension', '')}",
                ]
            )
        )
        supporting = [str(item).strip() for item in selected_frame.get("supporting_findings", []) if str(item).strip()]
        if supporting:
            lines.extend(["", "### Supporting Findings", *_markdown_list(supporting[:6])])
    elif row.frame_options:
        lines.extend(_markdown_list(["Story frames are available but no final frame is selected yet."]))
    else:
        lines.extend(_markdown_list(["Story frames have not been generated yet."]))

    lines.extend(["", "## Claim Map"])
    if row.claim_map:
        for claim in row.claim_map[:8]:
            lines.append(f"- {claim.get('claim_text', '')}")
            lines.append(f"  - Risk level: {claim.get('risk_level', 'unknown')}")
            if claim.get("rcmxt_summary"):
                lines.append(f"  - Evidence: {claim.get('rcmxt_summary')}")
            sources = claim.get("supporting_sources", []) or []
            if sources:
                lines.append(f"  - Sources: {', '.join(str(src) for src in sources[:4])}")
    else:
        lines.extend(_markdown_list(["No claim-level evidence map is available yet."]))

    lines.extend(["", "## Reviewer Risks"])
    if row.reviewer_risks:
        for risk in row.reviewer_risks[:8]:
            lines.append(f"- [{risk.get('severity', 'medium')}] {risk.get('title', '')}")
            if risk.get("section"):
                lines.append(f"  - Section: {risk.get('section')}")
            if risk.get("detail"):
                lines.append(f"  - Detail: {risk.get('detail')}")
            if risk.get("evidence_basis"):
                lines.append(f"  - Evidence basis: {risk.get('evidence_basis')}")
    else:
        lines.extend(_markdown_list(["No reviewer-risk analysis is available yet."]))

    lines.extend(["", "## Submission Checks"])
    if row.integrity_flags:
        for flag in row.integrity_flags[:10]:
            origin = (
                "quick draft scan"
                if str(flag.get("generated_by", "")).startswith("manuscript_session:quick_check")
                else "full audit"
                if str(flag.get("generated_by", "")).startswith("W7:")
                else "linked"
            )
            lines.append(f"- [{flag.get('severity', 'warning')}] {flag.get('title', '')}")
            lines.append(f"  - Source: {origin}")
            if flag.get("category"):
                lines.append(f"  - Category: {flag.get('category')}")
            if flag.get("detail"):
                lines.append(f"  - Detail: {flag.get('detail')}")
            if flag.get("suggestion"):
                lines.append(f"  - Suggestion: {flag.get('suggestion')}")
    else:
        lines.extend(_markdown_list(["No submission-check findings are available yet."]))

    lines.extend(["", "## Derived Outline"])
    if row.outline:
        for section in row.outline:
            lines.extend(["", f"### {section.get('title', 'Outline Section')}"])
            bullets = [str(item).strip() for item in section.get("bullets", []) if str(item).strip()]
            lines.extend(_markdown_list(bullets[:8] or ["No bullets available."]))
    else:
        lines.extend(_markdown_list(["The outline is not ready yet."]))

    lines.extend(["", "## Visible Fallbacks And Limitations"])
    if row.fallback_flags:
        lines.extend(
            _markdown_list(
                [
                    f"{flag.get('stage', 'stage')}: {flag.get('detail', '')}"
                    for flag in row.fallback_flags
                    if isinstance(flag, dict)
                ]
            )
        )
    else:
        lines.extend(_markdown_list(["No visible fallback flags are currently attached to this session."]))

    lines.extend(
        [
            "",
            "## Stage Statuses",
            *[
                f"- {status.stage.replace('_', ' ')}: {status.status.replace('_', ' ')} — {status.detail}"
                for status in _export_stage_statuses(row)
            ],
        ]
    )

    return "\n".join(lines).strip() + "\n"


def _severity_tone(value: str) -> str:
    lowered = value.lower()
    if lowered in {"high", "critical", "error", "failed"}:
        return "danger"
    if lowered in {"medium", "warning", "partial", "waiting_human"}:
        return "warn"
    if lowered in {"running"}:
        return "info"
    if lowered in {"not_started"}:
        return "muted"
    return "ok"


def _html_badge(text: str, tone: str) -> str:
    return f"<span class=\"badge badge-{tone}\">{_esc(text)}</span>"


def _html_list(items: list[str], empty: str) -> str:
    bullets = [f"<li>{_esc(item)}</li>" for item in items if item]
    if not bullets:
        bullets = [f"<li>{_esc(empty)}</li>"]
    return f"<ul>{''.join(bullets)}</ul>"


def _export_stage_statuses(row: ManuscriptSession) -> list[ManuscriptStageStatus]:
    try:
        return _to_response(row).stage_statuses
    except Exception:
        return [
            ManuscriptStageStatus(
                stage="session",
                status="partial",
                detail=f"Export fell back to session metadata only. Current phase: {row.phase}. Completion: {row.completion_state}.",
                source_workflow_id=None,
            )
        ]


def _build_defense_brief_html(row: ManuscriptSession) -> str:
    stage_statuses = _export_stage_statuses(row)

    session_meta = [
        _html_badge(f"Phase: {row.phase.replace('_', ' ')}", _severity_tone(row.phase)),
        _html_badge(f"Completion: {row.completion_state.replace('_', ' ')}", _severity_tone(row.completion_state)),
    ]
    if row.target_journal:
        session_meta.append(_html_badge(f"Target journal: {row.target_journal}", "info"))

    linked_workflows = _html_list(
        [f"{template}: {workflow_id}" for template, workflow_id in sorted((row.linked_workflows or {}).items())],
        "No linked workflows yet.",
    )

    if isinstance(row.selected_frame, dict):
        selected_frame = row.selected_frame
        story_frame_html = _html_list(
            [
                f"Selected frame: {selected_frame.get('frame_id', '')}",
                f"Narrative type: {str(selected_frame.get('narrative_type', '')).replace('_', ' ')}",
                f"Target tier: {selected_frame.get('target_tier', '')}",
                f"Hook: {selected_frame.get('hook', '')}",
                f"Core claim: {selected_frame.get('core_claim', '')}",
                f"Central tension: {selected_frame.get('central_tension', '')}",
            ],
            "No story frame details available.",
        )
        supporting_html = _html_list(
            [str(item).strip() for item in selected_frame.get("supporting_findings", []) if str(item).strip()][:6],
            "No supporting findings listed.",
        )
        story_frame_html += f"<h3>Supporting Findings</h3>{supporting_html}"
    elif row.frame_options:
        story_frame_html = _html_list(
            ["Story frames are available but no final frame is selected yet."],
            "Story frames are available but no final frame is selected yet.",
        )
    else:
        story_frame_html = _html_list(
            ["Story frames have not been generated yet."],
            "Story frames have not been generated yet.",
        )

    claim_rows: list[str] = []
    for claim in row.claim_map[:8]:
        sources = ", ".join(str(src) for src in (claim.get("supporting_sources", []) or [])[:4])
        claim_rows.append(
            "".join(
                [
                    "<tr>",
                    f"<td>{_esc(str(claim.get('claim_text', '')))}</td>",
                    f"<td>{_html_badge(str(claim.get('risk_level', 'unknown')), _severity_tone(str(claim.get('risk_level', 'unknown'))))}</td>",
                    f"<td>{_esc(str(claim.get('rcmxt_summary', '') or 'No evidence summary available.'))}</td>",
                    f"<td>{_esc(sources or 'No linked sources.')}</td>",
                    "</tr>",
                ]
            )
        )
    claim_map_html = (
        "<table><thead><tr><th>Claim</th><th>Risk</th><th>Evidence</th><th>Sources</th></tr></thead>"
        f"<tbody>{''.join(claim_rows)}</tbody></table>"
        if claim_rows
        else "<p class=\"empty\">No claim-level evidence map is available yet.</p>"
    )

    reviewer_cards = []
    for risk in row.reviewer_risks[:8]:
        reviewer_cards.append(
            "".join(
                [
                    "<article class=\"item-card\">",
                    f"<div class=\"item-head\"><h3>{_esc(str(risk.get('title', 'Reviewer risk')))}</h3>{_html_badge(str(risk.get('severity', 'medium')), _severity_tone(str(risk.get('severity', 'medium'))))}</div>",
                    f"<p><strong>Section:</strong> {_esc(str(risk.get('section', 'General')))}</p>",
                    f"<p>{_esc(str(risk.get('detail', '')))}</p>",
                    (
                        f"<p class=\"evidence\"><strong>Evidence basis:</strong> {_esc(str(risk.get('evidence_basis', '')))}</p>"
                        if risk.get("evidence_basis")
                        else ""
                    ),
                    "</article>",
                ]
            )
        )
    reviewer_risks_html = "".join(reviewer_cards) or "<p class=\"empty\">No reviewer-risk analysis is available yet.</p>"

    integrity_cards = []
    for flag in row.integrity_flags[:10]:
        origin = (
            "quick draft scan"
            if str(flag.get("generated_by", "")).startswith("manuscript_session:quick_check")
            else "full audit"
            if str(flag.get("generated_by", "")).startswith("W7:")
            else "linked"
        )
        integrity_cards.append(
            "".join(
                [
                    "<article class=\"item-card\">",
                    f"<div class=\"item-head\"><h3>{_esc(str(flag.get('title', 'Submission check')))}</h3>{_html_badge(str(flag.get('severity', 'warning')), _severity_tone(str(flag.get('severity', 'warning'))))}</div>",
                    f"<p><strong>Source:</strong> {_esc(origin)}",
                    f" · <strong>Category:</strong> {_esc(str(flag.get('category', 'unknown')))}</p>",
                    f"<p>{_esc(str(flag.get('detail', '')))}</p>",
                    (
                        f"<p class=\"evidence\"><strong>Suggestion:</strong> {_esc(str(flag.get('suggestion', '')))}</p>"
                        if flag.get("suggestion")
                        else ""
                    ),
                    "</article>",
                ]
            )
        )
    integrity_html = "".join(integrity_cards) or "<p class=\"empty\">No submission-check findings are available yet.</p>"

    outline_sections = []
    for section in row.outline:
        outline_sections.append(
            f"<article class=\"item-card\"><h3>{_esc(str(section.get('title', 'Outline Section')))}</h3>{_html_list([str(item).strip() for item in section.get('bullets', []) if str(item).strip()][:8], 'No bullets available.')}</article>"
        )
    outline_html = "".join(outline_sections) or "<p class=\"empty\">The outline is not ready yet.</p>"

    fallback_html = _html_list(
        [
            f"{flag.get('stage', 'stage')}: {flag.get('detail', '')}"
            for flag in row.fallback_flags
            if isinstance(flag, dict)
        ],
        "No visible fallback flags are currently attached to this session.",
    )
    stage_html = _html_list(
        [
            f"{status.stage.replace('_', ' ')}: {status.status.replace('_', ' ')} - {status.detail}"
            for status in stage_statuses
        ],
        "No stage status is available yet.",
    )

    notes_html = (
        f"<section class=\"panel\"><h2>Notes</h2><p>{_esc(row.notes.strip())}</p></section>"
        if row.notes.strip()
        else ""
    )

    return f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>{_esc(row.title)} — Defense Brief</title>
<style>
  :root {{
    color-scheme: light;
    --ink: #172033;
    --muted: #556179;
    --line: #d5dbe7;
    --panel: #ffffff;
    --canvas: #f4f7fb;
    --accent: #0f766e;
    --danger: #c2410c;
    --warn: #a16207;
    --ok: #166534;
    --info: #1d4ed8;
    --muted-badge: #475569;
  }}
  * {{ box-sizing: border-box; }}
  body {{
    margin: 0;
    background: var(--canvas);
    color: var(--ink);
    font-family: "IBM Plex Sans", "Segoe UI", sans-serif;
    line-height: 1.55;
  }}
  main {{
    max-width: 980px;
    margin: 0 auto;
    padding: 32px 24px 56px;
  }}
  header {{
    background: linear-gradient(135deg, #0f172a, #0f766e);
    color: white;
    border-radius: 24px;
    padding: 28px 28px 24px;
    margin-bottom: 24px;
  }}
  h1, h2, h3, p {{ margin-top: 0; }}
  h1 {{ font-size: 2rem; margin-bottom: 10px; }}
  h2 {{ font-size: 1.15rem; margin-bottom: 12px; }}
  h3 {{ font-size: 1rem; margin-bottom: 10px; }}
  .lede {{ color: rgba(255,255,255,0.86); max-width: 70ch; }}
  .badge-row {{ display: flex; flex-wrap: wrap; gap: 8px; margin-top: 14px; }}
  .badge {{
    display: inline-flex;
    align-items: center;
    border-radius: 999px;
    padding: 5px 10px;
    font-size: 0.78rem;
    font-weight: 600;
    background: rgba(255,255,255,0.12);
    border: 1px solid rgba(255,255,255,0.18);
  }}
  .badge-danger {{ color: #fee2e2; background: rgba(127,29,29,0.28); border-color: rgba(254,202,202,0.2); }}
  .badge-warn {{ color: #fef3c7; background: rgba(146,64,14,0.24); border-color: rgba(253,230,138,0.2); }}
  .badge-ok {{ color: #dcfce7; background: rgba(20,83,45,0.28); border-color: rgba(187,247,208,0.2); }}
  .badge-info {{ color: #dbeafe; background: rgba(30,64,175,0.26); border-color: rgba(191,219,254,0.2); }}
  .badge-muted {{ color: #e2e8f0; background: rgba(71,85,105,0.32); border-color: rgba(226,232,240,0.16); }}
  .grid {{
    display: grid;
    grid-template-columns: repeat(2, minmax(0, 1fr));
    gap: 16px;
    margin-bottom: 16px;
  }}
  .panel {{
    background: var(--panel);
    border: 1px solid var(--line);
    border-radius: 18px;
    padding: 18px 18px 16px;
    box-shadow: 0 10px 24px rgba(15, 23, 42, 0.05);
    break-inside: avoid;
  }}
  .panel-wide {{ margin-bottom: 16px; }}
  .item-card {{
    border: 1px solid var(--line);
    border-radius: 14px;
    padding: 14px 14px 12px;
    background: #fbfdff;
    margin-bottom: 12px;
  }}
  .item-head {{
    display: flex;
    align-items: flex-start;
    justify-content: space-between;
    gap: 10px;
    margin-bottom: 8px;
  }}
  .evidence {{ color: var(--muted); }}
  .empty {{ color: var(--muted); }}
  ul {{ margin: 0; padding-left: 20px; }}
  li + li {{ margin-top: 6px; }}
  table {{ width: 100%; border-collapse: collapse; font-size: 0.95rem; }}
  th, td {{ border: 1px solid var(--line); padding: 10px 12px; vertical-align: top; text-align: left; }}
  th {{ background: #eef4fb; }}
  footer {{
    color: var(--muted);
    font-size: 0.88rem;
    margin-top: 18px;
    text-align: right;
  }}
  @media (max-width: 860px) {{
    main {{ padding: 18px 14px 30px; }}
    .grid {{ grid-template-columns: 1fr; }}
    .item-head {{ flex-direction: column; }}
  }}
  @media print {{
    body {{ background: white; }}
    main {{ max-width: none; padding: 0; }}
    header {{
      background: white !important;
      color: var(--ink);
      border: 1px solid var(--line);
    }}
    .badge {{
      color: var(--ink) !important;
      background: white !important;
      border-color: var(--line) !important;
    }}
    .panel, .item-card, header {{
      box-shadow: none;
      break-inside: avoid;
    }}
  }}
</style>
</head>
<body>
<main>
  <header>
    <h1>{_esc(row.title)}</h1>
    <p class="lede">{_esc(row.query.strip() or 'No manuscript objective provided.')}</p>
    <div class="badge-row">{''.join(session_meta)}</div>
  </header>

  <section class="grid">
    <article class="panel">
      <h2>Linked Workflows</h2>
      {linked_workflows}
    </article>
    <article class="panel">
      <h2>Story Frame</h2>
      {story_frame_html}
    </article>
  </section>

  {notes_html}

  <section class="panel panel-wide">
    <h2>Claim Map</h2>
    {claim_map_html}
  </section>

  <section class="panel panel-wide">
    <h2>Reviewer Risks</h2>
    {reviewer_risks_html}
  </section>

  <section class="panel panel-wide">
    <h2>Submission Checks</h2>
    {integrity_html}
  </section>

  <section class="panel panel-wide">
    <h2>Derived Outline</h2>
    {outline_html}
  </section>

  <section class="grid">
    <article class="panel">
      <h2>Visible Fallbacks And Limitations</h2>
      {fallback_html}
    </article>
    <article class="panel">
      <h2>Stage Statuses</h2>
      {stage_html}
    </article>
  </section>

  <footer>Generated by BioTeam-AI Manuscript Studio. Use the browser print dialog to save this brief as PDF.</footer>
</main>
<script>
window.addEventListener("load", () => {{
  window.setTimeout(() => window.print(), 250);
}});
</script>
</body>
</html>
"""


def _docx_bullet(document, text: str, style: str = "List Bullet") -> None:
    cleaned = text.strip()
    if cleaned:
        document.add_paragraph(cleaned, style=style)


def _build_defense_brief_docx_bytes(row: ManuscriptSession) -> bytes:
    from docx import Document
    from docx.shared import Inches, Pt

    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    normal_style = document.styles["Normal"]
    normal_style.font.name = "Aptos"
    normal_style.font.size = Pt(10.5)

    title = document.add_heading(row.title, level=0)
    title.alignment = 0
    objective = document.add_paragraph()
    objective.add_run("Manuscript objective: ").bold = True
    objective.add_run(row.query.strip() or "No manuscript objective provided.")

    metadata = document.add_paragraph()
    metadata.add_run("Phase: ").bold = True
    metadata.add_run(row.phase.replace("_", " "))
    metadata.add_run(" | Completion: ").bold = True
    metadata.add_run(row.completion_state.replace("_", " "))
    if row.target_journal:
        metadata.add_run(" | Target journal: ").bold = True
        metadata.add_run(row.target_journal)

    document.add_heading("Linked Workflows", level=1)
    if row.linked_workflows:
        for template, workflow_id in sorted((row.linked_workflows or {}).items()):
            _docx_bullet(document, f"{template}: {workflow_id}")
    else:
        document.add_paragraph("No linked workflows yet.")

    if row.notes.strip():
        document.add_heading("Notes", level=1)
        document.add_paragraph(row.notes.strip())

    document.add_heading("Story Frame", level=1)
    if isinstance(row.selected_frame, dict):
        selected_frame = row.selected_frame
        for detail in (
            f"Selected frame: {selected_frame.get('frame_id', '')}",
            f"Narrative type: {str(selected_frame.get('narrative_type', '')).replace('_', ' ')}",
            f"Target tier: {selected_frame.get('target_tier', '')}",
            f"Hook: {selected_frame.get('hook', '')}",
            f"Core claim: {selected_frame.get('core_claim', '')}",
            f"Central tension: {selected_frame.get('central_tension', '')}",
        ):
            _docx_bullet(document, detail)
        supporting = [str(item).strip() for item in selected_frame.get("supporting_findings", []) if str(item).strip()]
        if supporting:
            supporting_heading = document.add_paragraph()
            supporting_heading.add_run("Supporting findings:").bold = True
            for item in supporting[:6]:
                _docx_bullet(document, item)
    elif row.frame_options:
        document.add_paragraph("Story frames are available but no final frame is selected yet.")
    else:
        document.add_paragraph("Story frames have not been generated yet.")

    document.add_heading("Claim Map", level=1)
    if row.claim_map:
        claim_table = document.add_table(rows=1, cols=4)
        claim_table.style = "Table Grid"
        headers = claim_table.rows[0].cells
        headers[0].text = "Claim"
        headers[1].text = "Risk"
        headers[2].text = "Evidence"
        headers[3].text = "Sources"
        for claim in row.claim_map[:8]:
            cells = claim_table.add_row().cells
            cells[0].text = str(claim.get("claim_text", ""))
            cells[1].text = str(claim.get("risk_level", "unknown"))
            cells[2].text = str(claim.get("rcmxt_summary", "") or "No evidence summary available.")
            cells[3].text = ", ".join(str(src) for src in (claim.get("supporting_sources", []) or [])[:4]) or "No linked sources."
    else:
        document.add_paragraph("No claim-level evidence map is available yet.")

    document.add_heading("Reviewer Risks", level=1)
    if row.reviewer_risks:
        for risk in row.reviewer_risks[:8]:
            paragraph = document.add_paragraph()
            paragraph.add_run(f"[{risk.get('severity', 'medium')}] ").bold = True
            paragraph.add_run(str(risk.get("title", "Reviewer risk"))).bold = True
            detail = document.add_paragraph()
            detail.add_run("Section: ").bold = True
            detail.add_run(str(risk.get("section", "General")))
            document.add_paragraph(str(risk.get("detail", "")))
            if risk.get("evidence_basis"):
                evidence = document.add_paragraph()
                evidence.add_run("Evidence basis: ").bold = True
                evidence.add_run(str(risk.get("evidence_basis", "")))
    else:
        document.add_paragraph("No reviewer-risk analysis is available yet.")

    document.add_heading("Submission Checks", level=1)
    if row.integrity_flags:
        for flag in row.integrity_flags[:10]:
            origin = (
                "quick draft scan"
                if str(flag.get("generated_by", "")).startswith("manuscript_session:quick_check")
                else "full audit"
                if str(flag.get("generated_by", "")).startswith("W7:")
                else "linked"
            )
            paragraph = document.add_paragraph()
            paragraph.add_run(f"[{flag.get('severity', 'warning')}] ").bold = True
            paragraph.add_run(str(flag.get("title", "Submission check"))).bold = True
            meta = document.add_paragraph()
            meta.add_run("Source: ").bold = True
            meta.add_run(origin)
            meta.add_run(" | Category: ").bold = True
            meta.add_run(str(flag.get("category", "unknown")))
            document.add_paragraph(str(flag.get("detail", "")))
            if flag.get("suggestion"):
                suggestion = document.add_paragraph()
                suggestion.add_run("Suggestion: ").bold = True
                suggestion.add_run(str(flag.get("suggestion", "")))
    else:
        document.add_paragraph("No submission-check findings are available yet.")

    document.add_heading("Derived Outline", level=1)
    if row.outline:
        for section_data in row.outline:
            document.add_heading(str(section_data.get("title", "Outline Section")), level=2)
            bullets = [str(item).strip() for item in section_data.get("bullets", []) if str(item).strip()]
            if bullets:
                for bullet in bullets[:8]:
                    _docx_bullet(document, bullet)
            else:
                document.add_paragraph("No bullets available.")
    else:
        document.add_paragraph("The outline is not ready yet.")

    document.add_heading("Visible Fallbacks And Limitations", level=1)
    if row.fallback_flags:
        for flag in row.fallback_flags:
            if isinstance(flag, dict):
                _docx_bullet(document, f"{flag.get('stage', 'stage')}: {flag.get('detail', '')}")
    else:
        document.add_paragraph("No visible fallback flags are currently attached to this session.")

    document.add_heading("Stage Statuses", level=1)
    for status in _export_stage_statuses(row):
        _docx_bullet(document, f"{status.stage.replace('_', ' ')}: {status.status.replace('_', ' ')} - {status.detail}")

    footer_paragraph = document.sections[0].footer.paragraphs[0]
    footer_paragraph.text = "Generated by BioTeam-AI Manuscript Studio."

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _get_linked_story_workflow_or_409(row: ManuscriptSession) -> WorkflowInstance:
    workflow_id = dict(row.linked_workflows or {}).get("W11")
    if not workflow_id:
        raise HTTPException(status_code=409, detail="No linked W11 workflow is available for this manuscript session.")
    return _get_workflow_or_404(workflow_id)


async def _inject_note_and_resume_story_workflow(workflow_id: str, note: str) -> None:
    from app.api.v1.workflows import InterveneRequest, intervene

    await intervene(
        workflow_id,
        InterveneRequest(action="inject_note", note=note, note_action="FREE_TEXT"),
    )
    await intervene(workflow_id, InterveneRequest(action="resume"))


def _materialize_uploaded_reviewer_paper(manuscript_session_id: str, filename: str, content: bytes) -> str:
    suffix = Path(filename or "").suffix.lower()
    if suffix not in _REVIEWER_UPLOAD_SUFFIXES:
        raise HTTPException(status_code=400, detail="Reviewer paper upload must be a .pdf, .docx, or .doc file.")
    if len(content) > 50_000_000:
        raise HTTPException(status_code=400, detail="Reviewer paper upload exceeds the 50MB limit.")

    safe_stem = re.sub(r"[^A-Za-z0-9._-]+", "_", Path(filename).stem or "reviewer_paper").strip("._")
    if not safe_stem:
        safe_stem = "reviewer_paper"

    upload_dir = Path(tempfile.gettempdir()) / "bioteam_manuscript_uploads" / manuscript_session_id
    upload_dir.mkdir(parents=True, exist_ok=True)
    upload_path = upload_dir / f"{safe_stem}_{uuid4().hex[:8]}{suffix}"
    upload_path.write_bytes(content)
    return str(upload_path.resolve(strict=False))


@router.get("/sessions", response_model=list[ManuscriptSessionResponse])
async def list_sessions() -> list[ManuscriptSessionResponse]:
    with Session(db_engine) as session:
        rows = session.exec(
            select(ManuscriptSession).order_by(ManuscriptSession.updated_at.desc())
        ).all()
        for row in rows:
            session.expunge(row)
    return [_to_response(row) for row in rows]


@router.post("/sessions", response_model=ManuscriptSessionResponse)
async def create_session(request: CreateManuscriptSessionRequest) -> ManuscriptSessionResponse:
    row = ManuscriptSession(
        title=(request.title or request.query[:120]).strip(),
        query=request.query,
        notes=request.notes,
        draft_text=request.draft_text,
        target_journal=request.target_journal,
        key_papers=request.key_papers,
    )
    row = _sync_session_outputs(row)
    row = _save_session(row)
    return _to_response(row)


@router.get("/sessions/{manuscript_session_id}", response_model=ManuscriptSessionResponse)
async def get_session(manuscript_session_id: str) -> ManuscriptSessionResponse:
    row = _get_session_or_404(manuscript_session_id)
    row = _sync_session_outputs(row)
    row = _save_session(row)
    return _to_response(row)


@router.get("/sessions/{manuscript_session_id}/defense-brief", response_model=ManuscriptDefenseBriefResponse)
async def get_defense_brief(manuscript_session_id: str) -> ManuscriptDefenseBriefResponse:
    row = _get_session_or_404(manuscript_session_id)
    row = _sync_session_outputs(row)
    row = _save_session(row)
    return ManuscriptDefenseBriefResponse(
        session_id=row.id,
        title=row.title,
        filename=_export_filename(row),
        markdown=_build_defense_brief_markdown(row),
    )


@router.get("/sessions/{manuscript_session_id}/defense-brief/print", response_model=ManuscriptDefenseBriefPrintResponse)
async def get_defense_brief_print(manuscript_session_id: str) -> ManuscriptDefenseBriefPrintResponse:
    row = _get_session_or_404(manuscript_session_id)
    row = _sync_session_outputs(row)
    row = _save_session(row)
    return ManuscriptDefenseBriefPrintResponse(
        session_id=row.id,
        title=row.title,
        filename=_export_filename(row, "html"),
        html=_build_defense_brief_html(row),
    )


@router.get("/sessions/{manuscript_session_id}/defense-brief/docx", response_model=ManuscriptDefenseBriefDocxResponse)
async def get_defense_brief_docx(manuscript_session_id: str) -> ManuscriptDefenseBriefDocxResponse:
    row = _get_session_or_404(manuscript_session_id)
    row = _sync_session_outputs(row)
    row = _save_session(row)
    content = _build_defense_brief_docx_bytes(row)
    return ManuscriptDefenseBriefDocxResponse(
        session_id=row.id,
        title=row.title,
        filename=_export_filename(row, "docx"),
        content_base64=base64.b64encode(content).decode("ascii"),
    )


@router.post("/sessions/{manuscript_session_id}/link-workflow", response_model=ManuscriptSessionResponse)
async def link_workflow(manuscript_session_id: str, request: LinkWorkflowRequest) -> ManuscriptSessionResponse:
    row = _get_session_or_404(manuscript_session_id)
    workflow = _get_workflow_or_404(request.workflow_id)
    if workflow.template not in _SUPPORTED_LINK_TEMPLATES:
        raise HTTPException(
            status_code=422,
            detail=f"Workflow template {workflow.template} is not supported for Manuscript Studio linking.",
        )

    linked = dict(row.linked_workflows or {})
    linked[workflow.template] = workflow.id
    row.linked_workflows = linked
    row = _sync_session_outputs(row)
    row = _save_session(row)
    return _to_response(row)


@router.post("/sessions/{manuscript_session_id}/select-frame", response_model=ManuscriptSessionResponse)
async def select_frame(manuscript_session_id: str, request: SelectStoryFrameRequest) -> ManuscriptSessionResponse:
    row = _get_session_or_404(manuscript_session_id)
    row = _sync_session_outputs(row)
    selected = next(
        (frame for frame in row.frame_options if str(frame.get("frame_id", "")).upper() == request.frame_id.upper()),
        None,
    )
    if selected is None:
        raise HTTPException(status_code=422, detail=f"Story frame not found in session: {request.frame_id}")
    selected_frame_id = str(selected.get("frame_id"))

    w11 = _get_linked_story_workflow_or_409(row) if dict(row.linked_workflows or {}).get("W11") else None
    if w11 is not None:
        manifest = dict(w11.session_manifest or {})
        if w11.state == "WAITING_HUMAN" and str(manifest.get("w11_phase", "")).strip().lower() == "awaiting_selection":
            await _inject_note_and_resume_story_workflow(
                w11.id,
                f"selected_frame_id: {request.frame_id}",
            )

    row = _get_session_or_404(manuscript_session_id)
    row.selected_frame_id = selected_frame_id
    row = _sync_session_outputs(row)
    row = _save_session(row)
    return _to_response(row)


@router.post("/sessions/{manuscript_session_id}/refresh", response_model=ManuscriptSessionResponse)
async def refresh_session(manuscript_session_id: str) -> ManuscriptSessionResponse:
    row = _get_session_or_404(manuscript_session_id)
    row = _sync_session_outputs(row)
    row = _save_session(row)
    return _to_response(row)


@router.post("/sessions/{manuscript_session_id}/run-story-frames", response_model=ManuscriptSessionResponse)
async def run_story_frames(manuscript_session_id: str) -> ManuscriptSessionResponse:
    row = _get_session_or_404(manuscript_session_id)
    from app.api.v1.workflows import CreateWorkflowRequest, create_workflow

    prompt_parts = [row.query.strip()]
    if row.notes.strip():
        prompt_parts.append(f"Research context:\n{row.notes.strip()}")
    if row.target_journal:
        prompt_parts.append(f"Target journal: {row.target_journal}")
    query = "\n\n".join(part for part in prompt_parts if part)[:2000]

    response = await create_workflow(
        CreateWorkflowRequest(template="W11", query=query, budget=2.0)
    )
    linked = dict(row.linked_workflows or {})
    linked["W11"] = response.workflow_id
    row.linked_workflows = linked
    row = _sync_session_outputs(row)
    row = _save_session(row)
    return _to_response(row)


@router.post("/sessions/{manuscript_session_id}/run-reviewer-risks", response_model=ManuscriptSessionResponse)
async def run_reviewer_risks(
    manuscript_session_id: str,
    request: RunReviewerRisksRequest,
) -> ManuscriptSessionResponse:
    row = _get_session_or_404(manuscript_session_id)
    from app.api.v1.workflows import CreateWorkflowRequest, create_workflow

    prompt_parts = [
        f"Review this biology manuscript for likely reviewer concerns.\n\nCore objective:\n{row.query.strip()}",
    ]
    if row.notes.strip():
        prompt_parts.append(f"Manuscript context:\n{row.notes.strip()}")
    if row.target_journal:
        prompt_parts.append(f"Target journal: {row.target_journal}")
    query = "\n\n".join(part for part in prompt_parts if part)[:2000]

    try:
        workflow_request = CreateWorkflowRequest(
            template="W8",
            query=query,
            budget=3.0,
            pdf_path=request.pdf_path,
        )
    except ValidationError as exc:
        first_error = exc.errors(include_url=False)[0] if exc.errors(include_url=False) else None
        detail = first_error.get("msg") if isinstance(first_error, dict) else "Invalid reviewer-risk input."
        raise HTTPException(status_code=422, detail=detail) from exc

    response = await create_workflow(workflow_request)
    linked = dict(row.linked_workflows or {})
    linked["W8"] = response.workflow_id
    row.linked_workflows = linked
    row = _sync_session_outputs(row)
    row = _save_session(row)
    return _to_response(row)


@router.post("/sessions/{manuscript_session_id}/upload-reviewer-paper", response_model=ManuscriptSessionResponse)
async def upload_reviewer_paper(
    manuscript_session_id: str,
    file: UploadFile = File(...),
) -> ManuscriptSessionResponse:
    _get_session_or_404(manuscript_session_id)
    content = await file.read()
    saved_path = _materialize_uploaded_reviewer_paper(
        manuscript_session_id,
        file.filename or "reviewer_paper.pdf",
        content,
    )
    return await run_reviewer_risks(
        manuscript_session_id,
        RunReviewerRisksRequest(pdf_path=saved_path),
    )


@router.post("/sessions/{manuscript_session_id}/resume-reviewer-risks", response_model=ManuscriptSessionResponse)
async def resume_reviewer_risks(manuscript_session_id: str) -> ManuscriptSessionResponse:
    row = _get_session_or_404(manuscript_session_id)
    row = _sync_session_outputs(row)
    workflow_id = dict(row.linked_workflows or {}).get("W8")
    if not workflow_id:
        raise HTTPException(status_code=409, detail="No linked W8 workflow is available for this manuscript session.")

    w8 = _get_workflow_or_404(workflow_id)
    if w8.state != "WAITING_HUMAN":
        raise HTTPException(
            status_code=409,
            detail=f"The linked W8 workflow is not waiting for reviewer approval (state={w8.state}).",
        )

    from app.api.v1.resume import ResumeRequest, resume_workflow

    await resume_workflow(w8.id, ResumeRequest())

    row = _get_session_or_404(manuscript_session_id)
    row = _sync_session_outputs(row)
    row = _save_session(row)
    return _to_response(row)


@router.post("/sessions/{manuscript_session_id}/approve-story-scope", response_model=ManuscriptSessionResponse)
async def approve_story_scope(
    manuscript_session_id: str,
    request: ApproveStoryScopeRequest,
) -> ManuscriptSessionResponse:
    row = _get_session_or_404(manuscript_session_id)
    row = _sync_session_outputs(row)
    w11 = _get_linked_story_workflow_or_409(row)
    manifest = dict(w11.session_manifest or {})
    phase = str(manifest.get("w11_phase", "")).strip().lower()
    if w11.state != "WAITING_HUMAN" or phase != "awaiting_scope":
        raise HTTPException(
            status_code=409,
            detail="The linked W11 workflow is not currently waiting for scope approval.",
        )

    await _inject_note_and_resume_story_workflow(
        w11.id,
        f"target_tier: {request.target_tier}",
    )

    row = _get_session_or_404(manuscript_session_id)
    row = _sync_session_outputs(row)
    row = _save_session(row)
    return _to_response(row)


def _persist_manual_integrity_run(result: dict, output: AgentOutput, duration_ms: int) -> None:
    findings_list = result.get("findings", [])
    with Session(db_engine) as session:
        for item in findings_list:
            session.add(
                AuditFinding(
                    category=item.get("category", "unknown"),
                    severity=item.get("severity", "info"),
                    title=item.get("title", ""),
                    description=item.get("description", ""),
                    source_text=item.get("source_text", ""),
                    suggestion=item.get("suggestion", ""),
                    confidence=item.get("confidence", 0.8),
                    checker=item.get("checker", ""),
                    finding_metadata=item.get("metadata", {}),
                )
            )
        session.add(
            AuditRun(
                trigger="manual",
                total_findings=result.get("total_findings", len(findings_list)),
                findings_by_severity=result.get("findings_by_severity", {}),
                findings_by_category=result.get("findings_by_category", {}),
                overall_level=result.get("overall_level", "clean"),
                summary=output.summary or "Manual manuscript submission check",
                cost=output.cost,
                duration_ms=duration_ms,
            )
        )
        session.commit()


@router.post("/sessions/{manuscript_session_id}/run-submission-checks", response_model=ManuscriptSessionResponse)
async def run_submission_checks(manuscript_session_id: str) -> ManuscriptSessionResponse:
    if _auditor_agent is None:
        raise HTTPException(status_code=503, detail="DataIntegrityAuditorAgent not available")

    row = _get_session_or_404(manuscript_session_id)
    text_parts = [row.draft_text.strip(), row.notes.strip()]
    text = "\n\n".join(part for part in text_parts if part)
    if not text.strip():
        raise HTTPException(
            status_code=422,
            detail="Add draft text or notes to the manuscript session before running submission checks.",
        )

    started = datetime.now(timezone.utc)
    output = await _auditor_agent.quick_check(text)
    duration_ms = int((datetime.now(timezone.utc) - started).total_seconds() * 1000)
    result = output.output or {}
    _persist_manual_integrity_run(result, output, duration_ms)

    quick_flags = [
        ManuscriptIntegrityFlag(
            title=str(item.get("title") or item.get("category", "Integrity finding")),
            severity=str(item.get("severity", "warning")),
            category=str(item.get("category", "unknown")),
            detail=str(item.get("description") or item.get("source_text", "")),
            suggestion=str(item.get("suggestion", "")),
            status="open",
            generated_by="manuscript_session:quick_check",
        ).model_dump(mode="json")
        for item in result.get("findings", [])
        if isinstance(item, dict)
    ]
    row.integrity_flags = _merge_unique_items(
        quick_flags,
        [item for item in row.integrity_flags or [] if isinstance(item, dict)],
        ("title", "detail", "generated_by"),
    )
    row.fallback_flags = _merge_unique_items(
        [item for item in row.fallback_flags or [] if isinstance(item, dict)],
        [
            ManuscriptFallbackFlag(
                stage="submission_checks",
                detail="Submission checks were run directly on session draft text, not through a linked W7 workflow.",
                provenance="manual_quick_check",
            ).model_dump(mode="json")
        ],
        ("stage", "detail", "source_workflow_id"),
    )
    row = _sync_session_outputs(row)
    row = _save_session(row)
    return _to_response(row)


@router.post("/sessions/{manuscript_session_id}/run-full-submission-audit", response_model=ManuscriptSessionResponse)
async def run_full_submission_audit(manuscript_session_id: str) -> ManuscriptSessionResponse:
    row = _get_session_or_404(manuscript_session_id)
    from app.api.v1.workflows import CreateWorkflowRequest, create_workflow

    prompt_parts = [f"Run a full biology submission audit for this manuscript objective:\n{row.query.strip()}"]
    if row.draft_text.strip():
        prompt_parts.append(f"Draft excerpt:\n{row.draft_text.strip()}")
    if row.notes.strip():
        prompt_parts.append(f"Submission notes:\n{row.notes.strip()}")
    if row.target_journal:
        prompt_parts.append(f"Target journal: {row.target_journal}")
    query = "\n\n".join(part for part in prompt_parts if part).strip()
    if not query:
        raise HTTPException(
            status_code=422,
            detail="Add a research objective, draft text, or notes to the manuscript session before starting a full submission audit.",
        )

    response = await create_workflow(
        CreateWorkflowRequest(template="W7", query=query[:2000], budget=3.0)
    )
    linked = dict(row.linked_workflows or {})
    linked["W7"] = response.workflow_id
    row.linked_workflows = linked
    row = _sync_session_outputs(row)
    row = _save_session(row)
    return _to_response(row)
